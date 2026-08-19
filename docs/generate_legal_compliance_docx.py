#!/usr/bin/env python3
"""Generate NetForensiq_Legal_Compliance.docx — government rules & legal compliance."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1B3A5C")
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[1 + r_idx].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "EDF2F7")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def heading(doc, text, level, color=RGBColor(0x1B, 0x3A, 0x5C)):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color

def body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(text)
    else:
        p.add_run(text)
    for run in p.runs:
        run.font.size = Pt(10)

# ── Document ─────────────────────────────────────────────────────────────────
doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6)
for level in range(1, 4):
    doc.styles[f"Heading {level}"].font.name = "Calibri"
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ═══ TITLE PAGE ═══
for _ in range(5):
    doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("NetForensiq")
run.bold = True; run.font.size = Pt(36); run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Government Rules & Legal Compliance Reference")
run.font.size = Pt(18); run.font.color.rgb = RGBColor(0x4A, 0x6F, 0xA5)

doc.add_paragraph()
tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tag.add_run(
    "Every statute, rule, and government directive\n"
    "that NetForensiq implements, references, or is designed to comply with."
)
run.italic = True; run.font.size = Pt(13); run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

for _ in range(4):
    doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(
    "KANAD S.H.I.E.L.D. 2026\n"
    "Category 2 — Problem Statement #8\n"
    "i-Hub Gujarat, Ahmedabad\n\nAugust 2026"
)
run.font.size = Pt(11); run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ═══ PAGE 2 — THE THREE NEW CODES ═══
doc.add_page_break()
heading(doc, "1. The New Criminal Law Regime (In Force 1 July 2024)", 1)

body(doc,
    "India's criminal law underwent its most comprehensive overhaul since independence "
    "when three new codes replaced the colonial-era statutes. Every feature of NetForensiq "
    "is designed against this post-1 July 2024 legal framework."
)

add_table(doc,
    ["New Code", "Full Name", "Replaces", "Relevance to NetForensiq"],
    [
        ["BNS", "Bharatiya Nyaya Sanhita, 2023",
         "Indian Penal Code 1860",
         "§111 (organised crime incl. cybercrime), §113 (terrorist act) — substantive offences"],
        ["BNSS", "Bharatiya Nagarik Suraksha Sanhita, 2023",
         "CrPC 1973",
         "§94 (production summons for CDR/IPDR), §105/185 (mandatory A/V recording), "
         "§173 (e-FIR/Zero FIR), §176(3) (mandatory forensic visit for ≥7yr offences)"],
        ["BSA", "Bharatiya Sakshya Adhiniyam, 2023",
         "Indian Evidence Act 1872",
         "§63 (electronic evidence admissibility + certificate) — the single most "
         "load-bearing provision for this project"],
    ],
    col_widths=[2.0, 5.0, 4.0, 5.5],
)

# ═══ PAGE 3 — BSA §63 DEEP DIVE ═══
doc.add_page_break()
heading(doc, "2. BSA Section 63 — Electronic Evidence Admissibility", 1)

body(doc,
    "Section 63 of the Bharatiya Sakshya Adhiniyam 2023 governs when an electronic record "
    "is admissible without producing the originating device. It is the direct successor to "
    "the long-litigated Section 65B of the Indian Evidence Act 1872. NetForensiq's entire "
    "evidence layer is built against this provision."
)

heading(doc, "2.1 The Four Statutory Conditions — §63(2)", 2)
bullet(doc, " The computer was used regularly for the relevant activity during the period.", bold_prefix="(a)")
bullet(doc, " Information of that kind was regularly fed into the computer.", bold_prefix="(b)")
bullet(doc, " The computer was operating properly throughout (or any malfunction did not affect accuracy).", bold_prefix="(c)")
bullet(doc, " The information reproduces or derives from what was fed in the ordinary course.", bold_prefix="(d)")

heading(doc, "2.2 The Mandatory Certificate — §63(4) + THE SCHEDULE", 2)
body(doc,
    "Section 63(4) requires a certificate to accompany every electronic record submitted "
    "for admission. The certificate form is prescribed in THE SCHEDULE to the Act, "
    "cross-referenced by §63(4)(c). It has two parts:"
)
bullet(doc, " Filled by the person in charge of the device (operator/custodian). "
       "Declares lawful control, regular operation, and the hash value with algorithm used. "
       "Checkboxes: Owned / Maintained / Managed / Operated.",
       bold_prefix="Part A (Party) —")
bullet(doc, " Filled by an expert. States the hash value independently and countersigns. "
       "Both signatures are required conjunctively ('signed by a person in charge … and an expert').",
       bold_prefix="Part B (Expert) —")

heading(doc, "2.3 Hash Algorithms Named in the Statute", 2)
body(doc,
    "THE SCHEDULE explicitly names three hash algorithms as checkbox options: "
    "SHA1, SHA256, MD5, plus an open 'Other (Legally acceptable standard)' field. "
    "NetForensiq computes all three so every checkbox on the form carries a value. "
    "SHA-256 is the primary digest relied upon; SHA-1 and MD5 are both cryptographically "
    "broken for collision resistance (per NIST guidance) and are included solely because "
    "the statutory form names them."
)

heading(doc, "2.4 How NetForensiq Implements §63", 2)
add_table(doc,
    ["Statutory Requirement", "NetForensiq Implementation"],
    [
        ["Hash value + algorithm (Schedule Part A/B)",
         "SHA-256 + MD5 + SHA-1 computed at ingest in a single streaming pass; "
         "algorithm declared as a structured field"],
        ["Device identification (type, make/model, serial, IMEI/MAC)",
         "Stored as structured EvidenceRecord fields matching the Schedule's checkbox list verbatim"],
        ["Dual signature (person in charge + expert)",
         "Certificate PDF requires both Part A and Part B signatures; "
         "missing either triggers a 'DRAFT — NOT A VALID CERTIFICATE' watermark"],
        ["Certificate form per THE SCHEDULE",
         "PDF reproduces the Schedule's exact wording, field order, and tick-boxes "
         "from the bare Act (indiacode.nic.in)"],
        ["Fields the tool cannot know (e.g., device colour)",
         "Printed blank for completion in ink — never filled with plausible values, "
         "as that would forge a statutory declaration"],
        ["Integrity re-verification",
         "Re-hashes the sealed PCAP at export time and compares against stored digest; "
         "certificate generation refuses to run if the check fails"],
    ],
    col_widths=[6.0, 10.5],
)

# ═══ PAGE 4 — BNSS PROVISIONS ═══
doc.add_page_break()
heading(doc, "3. BNSS Provisions Relevant to NetForensiq", 1)

heading(doc, "3.1 Section 94 — Production Summons for Digital Evidence", 2)
body(doc,
    "Section 94 BNSS (replacing §91 CrPC) empowers an investigating officer to issue a "
    "written requisition to a TSP or bank for CDR, IPDR, CAF, or account records. "
    "NetForensiq's evidence intake is designed to receive artefacts obtained through this "
    "channel and immediately seal them with hashing and chain-of-custody logging."
)

heading(doc, "3.2 Sections 105 & 185 — Mandatory A/V Recording of Search & Seizure", 2)
body(doc,
    "BNSS §105 mandates audio-video recording of every search and seizure. Warrant-search "
    "recordings must reach the Magistrate within 48 hours; warrantless within 'without delay'. "
    "The official app for this is eSakshya (NIC). NetForensiq is designed to accept "
    "eSakshya-format evidence (timestamped video + GPS + officer ID + hash) as a first-class "
    "record type."
)

heading(doc, "3.3 Section 173 — e-FIR / Zero FIR", 2)
body(doc,
    "Section 173(1) BNSS gives statutory recognition to Zero FIR (any police station, "
    "regardless of jurisdiction) and permits reporting by electronic communication. "
    "Gujarat launched a 'Cyber Financial Fraud e-Zero FIR' service on 27 July 2026 "
    "(Dy CM Harsh Sanghavi, developed with I4C) — a 1930-helpline complaint now "
    "auto-generates an FIR."
)

heading(doc, "3.4 Section 176(3) — Mandatory Forensic Visit (≥7yr offences)", 2)
body(doc,
    "For offences punishable with 7+ years, the IO must cause a forensic expert to visit "
    "the crime scene and videograph the process. States have up to 5 years to build capacity. "
    "NetForensiq's evidence-chain model supports forensic-visit artefact intake."
)

# ═══ PAGE 5 — IT ACT, CERT-In, DPDP, TELECOM ═══
doc.add_page_break()
heading(doc, "4. Other Government Rules & Directives", 1)

heading(doc, "4.1 IT Act 2000 — Sections Still in Force", 2)
body(doc, "The IT Act was NOT repealed by BNS/BNSS/BSA. Key surviving sections:")
add_table(doc,
    ["Section", "Subject", "Relevance"],
    [
        ["§66C", "Identity theft", "Fraudulent use of passwords/digital signatures; up to 3 yrs + ₹1L fine"],
        ["§66D", "Cheating by personation", "Core provision for online impersonation fraud"],
        ["§67/67A", "Obscene/sexually explicit material", "Publishing/transmitting electronically"],
        ["§69", "Interception/monitoring", "Government power to intercept in interest of sovereignty/security"],
        ["§69A", "Blocking of websites", "Government power to block public access to information"],
        ["§79", "Intermediary safe harbour", "Platforms not liable if due diligence observed"],
    ],
    col_widths=[2.5, 5.0, 9.0],
)

heading(doc, "4.2 CERT-In Directions 2022 (§70B(6), IT Act)", 2)
bullet(doc, " Mandatory reporting of cyber incidents to CERT-In within 6 hours.",
       bold_prefix="6-Hour Reporting —")
bullet(doc, " All ICT system logs must be retained for 180 days, stored within India.",
       bold_prefix="180-Day Log Retention —")

heading(doc, "4.3 DPDP Act 2023 — Law Enforcement Exemption", 2)
body(doc,
    "Section 17 exempts processing necessary for prevention/detection/investigation/ "
    "prosecution of offences. Police investigative tools sit inside this exemption. "
    "However, any citizen-facing component (complaint portal, self-service lookup) "
    "requires full DPDP hygiene: consent notice, purpose limitation, retention policy, "
    "grievance officer contact."
)

heading(doc, "4.4 Telecommunications Act 2023 — Interception Framework", 2)
body(doc,
    "Section 20 replaced the colonial Telegraph Act provisions. Interception requires an "
    "order from authorised Central/State agencies on grounds of public emergency/safety/ "
    "national security. Review Committees can set aside orders and order destruction. "
    "CDR/metadata (§94 BNSS summons) and live interception (Telecom Act 2023 order) are "
    "categorically different powers — NetForensiq never conflates them."
)

# ═══ PAGE 6 — CASE LAW ═══
doc.add_page_break()
heading(doc, "5. Key Case Law on Electronic Evidence", 1)

add_table(doc,
    ["Case", "Court / Date", "Holding & Relevance"],
    [
        ["Anvar P.V. v. P.K. Basheer\n(2014) 10 SCC 473",
         "Supreme Court\n18 Sep 2014",
         "§65B(4) certificate is mandatory for secondary electronic evidence. "
         "Overruled the relaxation in Navjot Sandhu."],
        ["Arjun Panditrao Khotkar v.\nKailash Gorantyal\n(2020) 7 SCC 1",
         "Supreme Court\n(3-judge bench)\n14 Jul 2020",
         "Upheld Anvar P.V. — certificate is 'a condition precedent'. "
         "Cannot be supplemented through oral evidence."],
        ["Pooranmal v. State of\nRajasthan (2026 INSC 217)",
         "Supreme Court\n2026",
         "Reaffirmed mandatory §65B certification and strict chain-of-custody "
         "for CDR/FSL evidence in circumstantial murder trial."],
        ["Kshitijbhai Patel v.\nDilipbhai Kanani\nSCA 120/2023",
         "Gujarat HC\nJ. Doshi\n8 May 2026",
         "§65B(4)/§63(4) certificate is 'a condition precedent'. Trial court's "
         "order sending evidence to FSL without deciding certificate question "
         "first was 'a patent illegality'. Also states: primary electronic "
         "documents must arrive 'along with the hash value'."],
        ["Mani Roy v. State of H.P.",
         "HP High Court\n27 May 2025",
         "eSakshya recording without §63(c) certificate is not admissible. "
         "Bail granted."],
        ["Shadab v. State of U.P.",
         "Allahabad HC\n5 Jan 2026",
         "Failure to videograph seizure per BNSS §105 showed 'negligence and "
         "arbitrariness'. Directed DGP to mandate eSakshya SOP."],
        ["Suresh v. State of Kerala",
         "Kerala HC\n26 Jul 2025",
         "Set aside murder conviction for 'irresponsible' crime-scene documentation. "
         "BNSS 'obligates' police to videotape every search/seizure."],
    ],
    col_widths=[4.0, 3.5, 9.0],
)

# ═══ PAGE 7 — CHAIN OF CUSTODY + COMPLIANCE CHECKLIST ═══
doc.add_page_break()
heading(doc, "6. Chain of Custody & Evidence Integrity Requirements", 1)

body(doc, "BSA §63(4)(c) and established forensic practice require:")
bullet(doc, " SHA-256 (or stated algorithm) at seizure/imaging. Any mismatch proves tampering.",
       bold_prefix="Hash at Ingest —")
bullet(doc, " Append-only, hash-chained event log. Each entry digests its predecessor.",
       bold_prefix="Immutable Audit Log —")
bullet(doc, " Every transfer (seizure → transport → lab → court) logged with actor, timestamp, location.",
       bold_prefix="Custody Transfer Events —")
bullet(doc, " Analysis never touches the sealed original. Working copies only.",
       bold_prefix="Original vs Working Copy —")
bullet(doc, " Re-hash at export; flag any mismatch before certificate generation.",
       bold_prefix="Verification at Export —")

heading(doc, "7. Compliance Checklist Summary", 1)

add_table(doc,
    ["#", "Requirement", "Source", "NetForensiq Status"],
    [
        ["1", "Hash at ingestion (SHA-256 + algorithm declared)", "BSA §63 Schedule", "✅ Implemented"],
        ["2", "Hash re-verification on export", "BSA §63 + forensic practice", "✅ Implemented"],
        ["3", "Immutable append-only audit log", "ISO 27037 + forensic practice", "✅ Hash-chained"],
        ["4", "BSA §63 certificate (Part A + Part B)", "BSA §63(4) + THE SCHEDULE", "✅ PDF generated"],
        ["5", "Custody-transfer events modelled", "BSA §63 + ISO 27037", "✅ CustodyEvent model"],
        ["6", "Original vs working copy separation", "Forensic best practice", "✅ Sealed store"],
        ["7", "eSakshya-compatible evidence intake", "BNSS §105/185", "✅ Supported"],
        ["8", "No live interception claims", "Telecom Act 2023 §20", "✅ Metadata only"],
        ["9", "DPDP hygiene on citizen-facing surfaces", "DPDP Act 2023 §17", "✅ Exemption noted"],
        ["10", "Synthetic demo data only (no real PII)", "DPDP Act + IT Act §66C/66D", "✅ Enforced"],
        ["11", "CERT-In 6-hour reporting awareness", "CERT-In Directions 2022", "✅ Documented"],
        ["12", "Dual-signature enforcement on certificate", "BSA §63(4) Schedule", "✅ DRAFT watermark"],
    ],
    col_widths=[1.0, 6.5, 4.5, 4.5],
)

# ═══ FOOTER ═══
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\n— NetForensiq · KANAD S.H.I.E.L.D. 2026 · i-Hub Gujarat, Ahmedabad —")
run.italic = True; run.font.size = Pt(9); run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Save ──
out_path = os.path.join(os.path.dirname(__file__), "NetForensiq_Legal_Compliance.docx")
doc.save(out_path)
print(f"✓ Saved: {out_path}")
