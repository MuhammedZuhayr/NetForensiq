#!/usr/bin/env python3
"""Generate NetForensiq_Overview.docx — a 5-page project overview."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background shading of a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a formatted table with header row shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1B3A5C")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[1 + r_idx].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "EDF2F7")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def add_heading_with_color(doc, text, level, color_rgb=RGBColor(0x1B, 0x3A, 0x5C)):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = color_rgb
    return heading


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(text)
    else:
        p.add_run(text)
    for run in p.runs:
        run.font.size = Pt(10)
    return p


def add_page_break(doc):
    doc.add_page_break()


# ── document ─────────────────────────────────────────────────────────────────

doc = Document()

# -- Global style tweaks --
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(2)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Calibri"

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 1 — Title Page
# ═══════════════════════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("NetForensiq")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Network & Packet Forensics Platform")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x4A, 0x6F, 0xA5)

doc.add_paragraph()

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run(
    "The chain-of-custody layer that makes network evidence\n"
    "stand up in an Indian court."
)
run.italic = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

for _ in range(4):
    doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(
    "KANAD S.H.I.E.L.D. 2026\n"
    "Category 2 — Problem Statement #8: Cyber Crime Investigation System\n"
    "i-Hub Gujarat, Ahmedabad\n\n"
    "August 2026"
)
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 2 — What Is NetForensiq & The Problem
# ═══════════════════════════════════════════════════════════════════════════════
add_page_break(doc)

add_heading_with_color(doc, "1. What Is NetForensiq?", level=1)

add_body(doc,
    "NetForensiq is a network and packet forensics platform that bridges a critical gap "
    "in the Indian digital evidence ecosystem. Existing tools like eSakshya seal scene video, "
    "and CCTNS Property Registers track physical objects. But a packet capture — the raw "
    "network traffic recorded during a cyber crime investigation — is neither a scene to "
    "videograph nor a physical object to log into a malkhana."
)

add_body(doc,
    "Network evidence falls between the two systems, and nothing currently covers it. "
    "NetForensiq is not another packet analyser like Wireshark or Arkime. It is the legal "
    "admissibility layer for network evidence in the Indian judicial system."
)

add_heading_with_color(doc, "1.1 The Problem We Solve", level=2)

add_body(doc,
    "When a cyber crime investigator captures network traffic (a PCAP file), three questions "
    "immediately arise that existing tools cannot answer:"
)

add_bullet(doc, " How do you prove the capture hasn't been tampered with after seizure?",
           bold_prefix="Integrity:")
add_bullet(doc, " Who handled this evidence, when, and what did they do with it?",
           bold_prefix="Chain of Custody:")
add_bullet(doc, " Can this evidence be presented in court under the Bharatiya Sakshya "
           "Adhiniyam (BSA) 2023?",
           bold_prefix="Legal Admissibility:")

add_heading_with_color(doc, "1.2 Why It Matters Now", level=2)

add_body(doc,
    "The Gujarat State Judicial Academy (GSJA) launched a Master Trainer Programme in "
    "July 2026, training judges to expect digital evidence in a specific shape: a unique ID, "
    "a cryptographic hash, and a BSA Section 63 certificate. NetForensiq deliberately mirrors "
    "that shape for network evidence — not imitation, but landing in the mental model "
    "the local judiciary already has."
)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 3 — Architecture & How It Works
# ═══════════════════════════════════════════════════════════════════════════════
add_page_break(doc)

add_heading_with_color(doc, "2. System Architecture", level=1)

add_body(doc,
    "NetForensiq is a full-stack application with a Django backend and a React frontend, "
    "designed for deployment in air-gapped police environments."
)

add_styled_table(doc,
    ["Layer", "Technology", "Purpose"],
    [
        ["Frontend", "React 19 + Vite + Material UI", "Dashboard, Findings, Evidence, Login pages"],
        ["Backend", "Django 6.0 + Django REST Framework", "API server, business logic, PDF generation"],
        ["Packet Engine", "Scapy (Python)", "PCAP parsing, flow assembly, feature extraction"],
        ["Detection", "7 rule-based detectors", "C2 beacons, DNS tunnels, port scans, exfiltration"],
        ["Evidence Layer", "SHA-256 + MD5 hashing", "Integrity verification, chain of custody, certificates"],
        ["Database", "SQLite (default) / PostgreSQL", "Zero-config for demos; PostgreSQL for production"],
        ["Auth", "JWT + role-based access", "Officer roles, badge IDs, department, admin approval"],
    ],
    col_widths=[3.5, 5.0, 7.5],
)

add_heading_with_color(doc, "2.1 How It Works — Step by Step", level=2)

add_body(doc, "The platform operates in three stages:")

# Stage 1
p = doc.add_paragraph()
run = p.add_run("Stage 1 — Capture & Ingest: ")
run.bold = True
run.font.size = Pt(10.5)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
p.add_run(
    "A PCAP file is imported via a management command. The moment it enters the system, "
    "SHA-256 and MD5 hashes are computed in a single streaming pass. A sealed copy is stored "
    "in the evidence store and is never modified again. The first chain-of-custody event "
    "(ACQUIRED) is logged. Scapy then parses every packet, assembles network flows, and "
    "extracts forensic features like Shannon entropy, TLS SNI, JA3 hashes, DNS label lengths, "
    "and TCP flag patterns."
)

# Stage 2
p = doc.add_paragraph()
run = p.add_run("Stage 2 — Detection & Analysis: ")
run.bold = True
run.font.size = Pt(10.5)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
p.add_run(
    "Seven rule-based detection engines run over the extracted features. Each rule carries "
    "its threshold source — either from published research (RITA, Snort, RFC 1035) or explicitly "
    "marked as [OUR HEURISTIC]. Findings are presented to analysts for triage: confirm, dismiss, "
    "or escalate. Every decision records the reviewer's identity and reasoning."
)

# Stage 3
p = doc.add_paragraph()
run = p.add_run("Stage 3 — Certification & Export: ")
run.bold = True
run.font.size = Pt(10.5)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
p.add_run(
    "When evidence is ready for court, NetForensiq re-hashes the sealed PCAP and compares "
    "it against the stored digest. If the integrity check passes, a BSA Section 63 certificate "
    "PDF is generated — reproducing THE SCHEDULE to the Bharatiya Sakshya Adhiniyam 2023 "
    "verbatim. The certificate requires dual signatures (Part A: person in charge; Part B: "
    "expert countersignature). An unsigned certificate is watermarked 'DRAFT — NOT A VALID "
    "CERTIFICATE' on every page."
)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 4 — Detection Engine & Legal Compliance
# ═══════════════════════════════════════════════════════════════════════════════
add_page_break(doc)

add_heading_with_color(doc, "3. Detection Engine", level=1)

add_body(doc,
    "The detection philosophy is: rules first, model second. A police panel will ask "
    "\"why did it flag this?\" — a rule answers, a black box does not. Every finding "
    "carries its threshold value, its provenance source, and full evidence JSON so an "
    "analyst can verify the logic independently."
)

add_styled_table(doc,
    ["Rule ID", "Detects", "Threshold Source"],
    [
        ["C2_BEACON_PERIODIC", "Command & Control beacons with regular timing",
         "RITA analyzer.go — MADM / median interval"],
        ["C2_BEACON_KEEPALIVE", "Periodic traffic inside a single persistent session",
         "Same MADM formula, applied intra-connection"],
        ["COVERT_CHANNEL_UNKNOWN_PORT", "Sustained egress to non-well-known port, no TLS SNI",
         "[OUR HEURISTIC]"],
        ["DNS_TUNNEL_LONG_LABEL", "Subdomain labels longer than 52 characters",
         "RFC 1035 §2.3.4; dnscat2; iodine"],
        ["RECON_PORT_SCAN", "One source probing many ports on one host",
         "Snort 3 port_scan (ports=25)"],
        ["EXFIL_VOLUME_ASYMMETRY", "Outbound volume exceeding p95 for the capture",
         "Relative threshold, floor at 100 KB"],
        ["ICMP_TUNNEL_OVERSIZED", "Oversized ICMP echo payloads in a sustained stream",
         "ping(8) baseline + [OUR HEURISTIC]"],
    ],
    col_widths=[4.5, 5.5, 5.5],
)

add_heading_with_color(doc, "3.1 Real-Traffic Validation", level=2)

add_body(doc,
    "The detection engine was validated against two real-world captures from "
    "malware-traffic-analysis.net (neither produced by us):"
)

add_styled_table(doc,
    ["Capture", "Purpose", "Result"],
    [
        ["AsyncRAT + XWorm infection\n(44 MB, 46k packets)",
         "True positive test",
         "Found 5 of 7 documented C2 flows, 0 false positives"],
        ["One week of server scans\n(28 MB, 362k packets)",
         "False positive test",
         "Reduced from 7,052 to 307 alerts after fixes"],
    ],
    col_widths=[5.0, 4.0, 7.0],
)

add_heading_with_color(doc, "4. Legal Compliance — BSA Section 63", level=1)

add_body(doc,
    "The Bharatiya Sakshya Adhiniyam 2023, Section 63, governs the admissibility of "
    "electronic records. NetForensiq's evidence layer implements:"
)

add_bullet(doc, " SHA-256 + MD5 dual hashing at ingest (as named in THE SCHEDULE itself)",
           bold_prefix="Cryptographic Integrity —")
add_bullet(doc, " Every access is logged as a CustodyEvent; each event digests its "
           "predecessor so altering any past entry breaks every subsequent link",
           bold_prefix="Hash-Chained Custody —")
add_bullet(doc, " PDF reproduces THE SCHEDULE verbatim with the same wording, field order, "
           "and tick-boxes from the bare Act",
           bold_prefix="Certificate Generation —")
add_bullet(doc, " Where NetForensiq does not hold a fact the Schedule asks for (e.g., "
           "device colour, parent's name), the line is printed blank for a human to complete "
           "in ink. Filling it with plausible values would be forging a statutory declaration.",
           bold_prefix="Statutory Blanks —")
add_bullet(doc, " Certificate requires both Part A (person in charge) and Part B (expert). "
           "Missing either triggers a DRAFT watermark on every page.",
           bold_prefix="Dual Signature —")


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 5 — Tech Stack, Testing, and Summary
# ═══════════════════════════════════════════════════════════════════════════════
add_page_break(doc)

add_heading_with_color(doc, "5. Technology Stack", level=1)

add_styled_table(doc,
    ["Component", "Technology", "Version"],
    [
        ["Language (Backend)", "Python", "3.10+"],
        ["Web Framework", "Django", "6.0"],
        ["REST API", "Django REST Framework", "Latest"],
        ["Packet Parsing", "Scapy", "Latest"],
        ["Language (Frontend)", "JavaScript (ES2020+)", "—"],
        ["UI Framework", "React", "19"],
        ["Build Tool", "Vite", "Latest"],
        ["UI Components", "Material UI (MUI)", "Latest"],
        ["E2E Testing", "Playwright", "Latest"],
        ["Authentication", "JWT (SimpleJWT)", "Latest"],
        ["Database", "SQLite / PostgreSQL", "3.x / 15+"],
    ],
    col_widths=[4.5, 5.5, 3.0],
)

add_heading_with_color(doc, "5.1 Test Coverage", level=2)

add_styled_table(doc,
    ["Suite", "Count", "What It Covers"],
    [
        ["Backend Unit Tests", "61", "Feature maths, timestamp fidelity, all attack types, "
         "false-positive guards, DNS aggregation, threshold provenance, IPv6, hashing, "
         "tamper detection, custody-chain breakage, certificate refusal on failed integrity"],
        ["Playwright E2E Tests", "15", "Auth guard, dashboard figures matching the API, "
         "absence of placeholder strings, threshold inspection, triage round-trip, "
         "custody verdict, certificate download"],
    ],
    col_widths=[4.5, 2.0, 9.5],
)

add_heading_with_color(doc, "5.2 Project Structure", level=2)

add_body(doc, "The repository is organized as follows:")

add_styled_table(doc,
    ["Directory", "Contents"],
    [
        ["backend/accounts/", "Authentication, roles, badge IDs, audit logging"],
        ["backend/capture/", "Packet processor, flow assembly, detection engine, synthetic generator"],
        ["backend/evidence/", "Evidence integrity, chain of custody, BSA s.63 certificate PDF"],
        ["frontend/src/pages/", "Dashboard, Findings, Evidence, Login pages"],
        ["frontend/src/services/", "API client modules (forensics.js, auth.js)"],
        ["frontend/e2e/", "Playwright end-to-end test suite"],
        ["research/", "19 research documents (legal, technical, intelligence)"],
        ["scripts/", "Utility and verification scripts"],
    ],
    col_widths=[5.0, 11.0],
)

add_heading_with_color(doc, "6. Summary", level=1)

add_body(doc,
    "NetForensiq addresses a specific, unmet need: making network evidence legally admissible "
    "in Indian courts. It combines a Scapy-based packet engine, a rule-based detection system "
    "with cited thresholds, and a BSA Section 63 certificate generator — all in one platform. "
    "The system is designed for air-gapped police environments, requires zero external services "
    "to function, and has been validated against real-world malware traffic."
)

add_body(doc,
    "The project stands on three principles: (1) no hardcoded demo data — every number on "
    "screen comes from the database; (2) no invented thresholds — every detection parameter "
    "carries a citation or is explicitly labelled [OUR HEURISTIC]; and (3) no overclaiming — "
    "synthetic-data performance is reported as such."
)

# ── Footer ────────────────────────────────────────────────────────────────────

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\n— NetForensiq · KANAD S.H.I.E.L.D. 2026 · i-Hub Gujarat, Ahmedabad —")
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


# ── Save ──────────────────────────────────────────────────────────────────────

out_path = os.path.join(os.path.dirname(__file__), "NetForensiq_Overview.docx")
doc.save(out_path)
print(f"✓ Saved: {out_path}")
